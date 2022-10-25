# coding=utf-8
"""
modified by : Guttappa sajjan
Performs face detection and recognition in realtime.

Based on code from https://github.com/shanren7/real_time_face_recognition
"""
import argparse
# import playsound
import sys
import time
import datetime
import cv2
import os
import face
from datetime import date
from datetime import datetime
import pandas as pd

from PIL import Image
import csv
import errno
import xlsxwriter
import numpy as np
import imutils
from imutils.video import VideoStream
from tensorflow.keras.preprocessing import image
import os
from PIL import ImageDraw, ImageFont

from vidgear.gears import VideoGear
import cv2
import time


from flask import Flask, render_template, Response
import cv2

import os
# Import docx NOT python-docx
import docx

####### Telegram Integration #############3
import telepot


token = '5700004463:AAHnk8J83V1zBNis9YgmrdoNqY2C4t4mC4I' # telegram token
receiver_id = 754000023# https://api.telegram.org/bot<TOKEN>/getUpdates

##########
# Create an instance of a word document
doc = docx.Document()
# account_sid = os.environ['AC8b74322ab11b2c3cf473e6c55cdb1291']
# auth_token = os.environ['5f236fc5d2254b9fc5b3c9726b59329a']
# Add a Title to the document

# from twilio.rest import Client
# client = Client(account_sid, auth_token)
#if datetime.datetime.now().time() < datetime.time(hour = 13, minute = 35, second =1):

# timeStamp = time.strftime("%-I:%M %p")
# # #print('..............',timeStamp)
st_main = str(date.today())
# if datetime.datetime.now().time() < datetime.time(hour = 13, minute = 35, second =1):
#     #print("chekins")
#     doc.add_heading('Report '+ st_main + "--(Check-In)", 0)
# else:
#     #print('checkouts')
#     doc.add_heading('Report '+ st_main + "--(Check-Out)", 0)
PEOPLE_FOLDER = os.path.join('Attendance', st_main)

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = PEOPLE_FOLDER
global path
path = ''
#camera = cv2.VideoCapture('rtsp://freja.hiof.no:1935/rtplive/_definst_/hessdalen03.stream')  # use 0 for web camera
#  for cctv camera use rtsp://username:password@ip_address:554/user=username_password='password'_channel=channel_number_stream=0.sdp' instead of camera
# for local webcam use cv2.VideoCapture(0)

# def gen_frames():  # generate frame by frame from camera
#     # while True:
#     #     # # Capture frame-by-frame
#     #     # success, frame = camera.read()  # read the camera frame
#     #     # if not success:
#     #     #     break
#     #     # else:
#     #     #     ret, buffer = cv2.imencode('.jpg', frame)
#     frame = buffer.tobytes()
#     yield (b'--frame\r\n'
#            b'Content-Type: image/jpeg\r\n\r\n' + frame + b'\r\n')  # concat frame one by one and show result

# def sound_alarm(path):
#     # play an alarm sound
#     playsound.playsound(path)
#######
os.environ["CUDA_DEVICE_ORDER"]="PCI_BUS_ID"  
os.environ["CUDA_VISIBLE_DEVICES"]="0"
ts = time.time()
# st_main = str(date.today())
# print(st_main)
# print("i came here once..")
# #dir_name='Attendance'
# # creating a folder named data
# os.chdir('Attendance/')
# try:
#     if not os.path.exists(st_main):
#         os.makedirs(st_main)
#         ### cd to the folder
#         os.chdir(st_main)
#     else:
#         #os.makedirs('Attendance/'+st_main)
#         ### cd to the folder
#         os.chdir(st_main)

# except OSError as e:
#     if e.errno != errno.EEXIST:
#         raise   
#     # time.sleep might help here
#     pass
########
# Creating a table object
table = doc.add_table(rows=1, cols=4)
# Adding heading in the 1st row of the table
row = table.rows[0].cells
row[0].text = 'Employees_Name'
row[1].text = 'Login_Date'
row[2].text = 'Login_Time'
row[3].text = 'Image'
def add_name(img,faces):
    if faces is not None:
        for face in faces:

            if face.name is not None:
                # nn= face.name
                cv2.putText(img, "Detected Person : "+face.name, (30, 30),cv2.FONT_HERSHEY_SIMPLEX, 0.8, (0, 255, 0), 2)

########
def store_single_disk(image, label):
    """ 
    @author : Guttappa Sajjan
    Stores a single image as a .png file on disk.
        Parameters:
        ---------------
        image       image array, (32, 32, 3) to be stored
        label        image label
    """
    ts = time.time()
    st = datetime.fromtimestamp(ts).strftime('%d-%m-%Y_%H-%M-%S')
    Image.fromarray(image).save(f"{label}_{st}.png")
    return f"{label}_{st}.png"

#########3
def mark_attendance(csvData):
    with open('attendance.csv', 'a') as csvFile:
        writer = csv.writer(csvFile)
        writer.writerows(csvData)
        rows = open('attendance.csv').read().split('\n')                                               
        newrows = []
        for row in rows:
            if row not in newrows:
                newrows.append(row)

        csvFile = open('FinalAttendance.csv', 'w')
        csvFile.write('\n'.join(newrows))
        #csvFile.to_excel ('FinalAttendance.xlsx', index = None, header=True)
        # df = pd.read_csv('FinalAttendance.csv')
        # df.to_excel('FinalAttendance.xlsx', sheet_name='gkz', index=True)  # index=True to write row index
        csvFile.close()

#########            
def add_overlays(frame, faces, frame_rate,today_date):
    #draw = ImageDraw.Draw(img)
    #fs = max(20, round(frame.shape[0] * frame.shape[1] * 0.000005))
    # print('fsssssssssss',fs)
    #font = ImageFont.truetype('fonts/font.ttf', fs)
    margin = 5
    #ggs = 0
    #print(ggs)
    if faces is not None:
        #main_dict={}
        #draw = ImageDraw.Draw(img)
        #fs = max(20, round(frame.shape[0] * frame.shape[1] * 0.000005))
        # print('fsssssssssss',fs)
        #font = ImageFont.truetype('fonts/font.ttf', fs)
        margin = 5
        for face in faces:
            face_bb = face.bounding_box.astype(int)
            (startX, startY, endX, endY)=face.bounding_box.astype(int)
            #x, y, w, h = [ v for v in face_bb]
            #top, right, bottom, left= [ v for v in face_bb]
            #print(top, right, bottom, left)
            #cv2.rectangle(img, (x,y), (x+w,y+h), (255,255,255))
            # x0,y0=(int(x),int(y))
            # x1,y1=(int(x+w),int(y+h))
            # print(x0,y0)
            # print(x1,y1)
            sub_face = frame[startY:endY, startX:endX]
            target_size=(320, 320)
            #if grayscale == True:
            sub_face = cv2.cvtColor(sub_face, cv2.COLOR_BGR2GRAY)
                
            crop_face = cv2.resize(sub_face, target_size)
            global path
            Image.fromarray(crop_face).save(f"{face.name}.png")
            path = f"{face.name}.png"
            nanu=store_single_disk(crop_face, face.name)
            #print('pathjfdhdsfhdskjhdskfjhdskjfhsdkjf',path)
            # if face_bb[2]>130:
            cv2.rectangle(frame,
                      (face_bb[0], face_bb[1]), (face_bb[2], face_bb[3]),
                      (0, 255, 0), 2)


            if face.name is not None:

                text = face.name.upper()
                #print('fsssssssssss',text)
                #text_size = font.getsize(text)
                cv2.putText(frame, text, (face_bb[0], face_bb[3]),
                            cv2.FONT_HERSHEY_SIMPLEX, 1, (255, 0, 255),
                            thickness=2, lineType=2)
                # message = client.messages \
                # .create(
                #      body='Detected face is : '+text,
                #      from_ =  +19706477495,
                #      to = +918147602693
                #  )

                # print(message.sid)
                date = datetime.fromtimestamp(ts).strftime('%Y-%m-%d')
                #timeStamp = datetime.now().strftime('%H:%M:%S')
                timeStamp = time.strftime("%-I:%M %p") ##%r for seconds and IST print

                #attendance.loc[len(attendance)] = [face.name, date, timeStamp]
                #print(attendance)

            #crop_face = frame[face_bb[1]:face_bb[1]+face_bb[3], face_bb[0]:face_bb[0]+face_bb[2]]
            list_of_tuples = list(zip(face.name, today_date,timeStamp))
             
            # Assign data to tuples.
            

            csvData = [['Employees_Name', 'Login_Date','Login_Time'], [face.name, today_date,timeStamp]]
            # Adding a row and then adding data in it.
            # initialize data of lists.
            data = {'Employees_Name': face.name,
                    'Login_Date': today_date,'Login_Time':timeStamp}
            print(data)
            df = pd.DataFrame(list_of_tuples,columns=['Employees_Name', 'Login_Date','Login_Time'])
            ###############
            if face.name != 'unknown':
                #print("i became : ", ggs)
                bot = telepot.Bot(token)
                bot.sendMessage(receiver_id, 'Detected face: '+face.name) # send a activation message to telegram receiver id
                bot.sendPhoto(receiver_id, photo=open(nanu, 'rb')) # send message to telegram

        ###########
            #     print("iam not unknown..")
            #     row = table.add_row().cells
            #     # Converting id to string as table can only take string input
            #     row[0].text = face.name
            #     row[1].text = today_date
            #     row[2].text = timeStamp
            #     #row[1].text = 

            #     paragraph = row[3].paragraphs[0]
            #     run = paragraph.add_run()
            #     run.add_picture(path, width = 1400000, height = 1400000)
            #     #ggs+=1
            #     if ggs ==4:
            #         break
                

            timeStamp = time.strftime("%-I:%M %p")
            # #print('..............',timeStamp)
            #st_main = str(date.today())
            mark_attendance(csvData)
            ##### added on 14-10-2022#######################
            if timeStamp < '1:00 PM':
                #print("chekins")
                try:
                    if not os.path.exists('CheckIns'):
                        os.makedirs('CheckIns')
                        ### cd to the folder
                        os.chdir('CheckIns')
                    else:
                        #os.makedirs('Attendance/'+st_main)
                        ### cd to the folder
                        os.chdir('CheckIns')

                except OSError as e:
                    if e.errno != errno.EEXIST:
                        raise   
                    # time.sleep might help here
                    pass
                mark_attendance(csvData)
                #doc.add_heading('Report '+ st_main + "--(Check-In)", 0)
            else:
                try:
                    if not os.path.exists('CheckOuts'):
                        os.makedirs('CheckOuts')
                        ### cd to the folder
                        os.chdir('CheckOuts')
                    else:
                        #os.makedirs('Attendance/'+st_main)
                        ### cd to the folder
                        os.chdir('CheckOuts')

                except OSError as e:
                    if e.errno != errno.EEXIST:
                        raise   
                    # time.sleep might help here
                    pass
        ########################## ends here ##################################
            #     #os.makedirs('CheckOuts')
            #     #print('checkouts')
                # mark_attendance(csvData)
                #doc.add_heading('Report '+ st_main + "--(Check-Out)", 0)
            
            # start=time.time()
            # grayscale = True
            # target_size=(320, 320)
            #if grayscale == True:
            #sub_face = cv2.cvtColor(sub_face, cv2.COLOR_BGR2GRAY)
                
            #crop_face = cv2.resize(sub_face, target_size)
            #img_pixels = image.img_to_array(crop_face)
            #img_pixels = np.expand_dims(img_pixels, axis = 0)
            #img_pixels /= 255 #normalize input in [0, 1]
            # count=0
            # if (time.time() - start > 3):
                # count+=1
            #ggs= True
        
    # Now save the document to a location
    

video_capture = VideoStream('rtsp://admin:Ashlesha123@192.168.0.170').start()
#video_capture = VideoStream(0).start()
#video_capture= cv2.VideoCapture(-1)
def gen_frames():
#def main(args):
    frame_interval = 3  # Number of frames after which to run face detection
    fps_display_interval = 5  # seconds
    frame_rate = 0
    frame_count = 0
    sleep_time = 2
    num_retries = 2
    
    # st_main = str(date.today())
    # timeStamp = time.strftime("%-I:%M %p")
    # #print('..............',timeStamp)
    # #st_main = str(date.today())
    # if timeStamp < '1:00 PM':
    #     print("chekins")
    #     doc.add_heading('Report '+ st_main + "--(Check-In)", 0)
    # else:
    #     print('checkouts')
    #     doc.add_heading('Report '+ st_main + "--(Check-Out)", 0)
    # for x in range(0, num_retries):  
    #     try:

    #         #video_capture = cv2.VideoCapture('/home/guttappa/Desktop/Guttappa/filename.avi')
    #         #print("i came in try again ..")
    #         video_capture = VideoStream('rtsp://admin:Ashlesha123@192.168.0.170').start()
    #         #print("i came here again trying ..")
    #         #video_capture = VideoStream(0).start()
    #         # face_recognition = face.Recognition()
    #         # start_time = time.time()
    #         video = video_capture.read()
    #         #####
            
    #         #frameA=cv2.resize(video,(1024 ,768))
    #         frame = imutils.resize(video, width=1100)
    #         #r = cv2.selectROI(("ROI",video))
    #         (x,y,w,h) = cv2.selectROI("Select ROI", frame)
    #         cv2.destroyAllWindows()
            # define and start the stream on first source ( For e.g #0 index device)
            #stream1 = VideoGear(source=0, logging=True).start() 

            # define and start the stream on second source ( For e.g #1 index device)
            #stream2 = VideoGear(source='rtsp://admin:Ashlesha123@192.168.0.170', logging=True).start() 

    face_recognition = face.Recognition()
    start_time = time.time()

    ######

    # if args.debug:
    #     print("Debug enabled")
    #     face.debug = True
    ggs = 0
    while True:
        
        ts = time.time()
        st_main = str(date.today())
        #print(st_main)
        #print("i came in while once..")
        #dir_name='Attendance'
        # creating a folder named data

        os.chdir('/home/prixgen-gpu/Desktop/FaceGen/facenet/contributed/Attendance')
        try:
            if not os.path.exists(st_main):
                os.makedirs(st_main)
                ### cd to the folder
                os.chdir(st_main)
            else:
                #os.makedirs('Attendance/'+st_main)
                ### cd to the folder
                os.chdir(st_main)

        except OSError as e:
            if e.errno != errno.EEXIST:
                raise   
            # time.sleep might help here
            pass
        # Capture frame-by-frame
        img = video_capture.read()
        #print(img)
        #img = imutils.resize(img, width=1100)
        #frame=img[y:y+h, x:x+w]
        #frame=img[int(r[1]):int(r[1]+r[3]), int(r[0]):int(r[0]+r[2])]


        ###

        #frameA = stream1.read()
        # read frames from stream1

        #frameB = stream2.read()
        #frameB= cv2.cvtColor(frame, cv2.COLOR_BGR2LAB)
        alpha =2.0  # Contrast control (1.0-3.0)
        beta = 0 # Brightness control (0-100)

        #adjusted = cv2.convertScaleAbs(frame, alpha=alpha, beta=beta)

        # ALARM_ON = False
        #frame = cv2.cvtColor(frame, cv2.COLOR_BGR2GRAY)
        #print('fsssssssssss',adjusted.size[1])
        ###

        # scale_percent = 50 # percent of original size
        # width = int(frame.shape[1] * scale_percent / 100)
        # height = int(frame.shape[0] * scale_percent / 100)
        # dim = (width, height)
          
        # # # resize image
        # resized = cv2.resize(frame, dim, interpolation = cv2.INTER_AREA)

        # ###frame2
        # scale_percent = 40 # percent of original size
        # width = int(frameB.shape[1] * scale_percent / 100)
        # height = int(frameB.shape[0] * scale_percent / 100)
        # dim = (width, height)
          
        # # resize image
        # resizedB = cv2.resize(frameB, dim, interpolation = cv2.INTER_AREA)



        #frameA=cv2.resize(frameA,(620,480))
        #frame=cv2.resize(frame,(620,480))
        # height, width, channels = frame.shape
        # upper_left = (width // 4, height // 4)
        # bottom_right = (width * 4 // 4, height * 4 // 4)

        #cv2.rectangle(frame,upper_left, bottom_right, (100, 50, 200), 5)
        #rect_img = frame[upper_left[1]: bottom_right[1] + 1, upper_left[0]: bottom_right[0] + 1]
        #frame =cv2.imread('/home/guttappa/Downloads/IMG-20210202-WA0008.jpg')

        if (frame_count % frame_interval) == 0:
            #facesA = face_recognition.identify(frameA)
            faces = face_recognition.identify(img)
            # if not ALARM_ON:
            #     ALARM_ON = True

            # if args["alarm"] != "":
            #     t = Thread(target=sound_alarm,
            #         args=(args["alarm"],))
            #     t.deamon = True
            #     t.start()


            # # Check our current fps
            # end_time = time.time()
            # if (end_time - start_time) > fps_display_interval:
            #     frame_rate = int(frame_count / (end_time - start_time))
            #     start_time = time.time()
            #     frame_count = 0
            
        #add_overlays(frameA, facesA, frame_rate)
        #print(faces)
        if faces:
            add_overlays(img, faces, frame_rate,st_main)
            #print("i came here for faces")
            #ggs+=1
        #print(faces.name)
        # store_single_disk(frame, image_id, label)

        frame_count += 1
        #add_name(frameA,facesA)
        add_name(img,faces)
        ####adding date and time to the frame
        #timestamp = datetime.now()
        # tm = timestamp.strftime("%A %d %B %Y %I:%M:%S%p")        
        #cv2.putText(resized, nn, (10, resized.shape[0] - 10),cv2.FONT_HERSHEY_SIMPLEX, 0.8, (0, 255, 0), 2)
        #cv2.imshow('Check-in stream...', frameA)
        #cv2.imshow('Check-in stream...', frame)
        ret, buffer = cv2.imencode('.jpg', img)
        img = buffer.tobytes()
        yield (b'--frame\r\n'
               b'Content-Type: image/jpeg\r\n\r\n' + img + b'\r\n')  # concat frame one by one and show result

        #cv2.imshow('Check-in adjusted...', img)
        # if cv2.waitKey(100) & 0xFF == ord('q'):
        #     break
        #doc.save('report_'+st_main+'.docx')
        

    # cv2.destroyAllWindows()
    # video_capture.stop()
        # except (OSError, TypeError, AttributeError,  SyntaxError, KeyboardInterrupt, cv2.error) as e:
        #         print(e)
        #         from time import sleep
        #         #logger.error(e)
        #         cv2.destroyAllWindows()
        #         video_capture.stop()
        #         sleep(20)
        #         pass



    # When everything is done, release the capture
    #video_capture.release()
    #print("Please check your camara and internet connection..! ")
    video_capture.stop()
    cv2.destroyAllWindows()

@app.route('/video_feed')
def video_feed():
    #Video streaming route. Put this in the src attribute of an img tag
    return Response(gen_frames(), mimetype='multipart/x-mixed-replace; boundary=frame')


@app.route('/')
def index():
    """Video streaming home page."""
    # full_filename = '/home/prixgen-gpu/Desktop/FaceGen/facenet/contributed/Attendance/2022-06-17/unknown_17-06-2022_18-04-02.jpg'
    # print("full_filename.......................",full_filename)
    return render_template('index.html')


@app.route('/requests',methods=['POST','GET'])
def tasks():
    global switch,camera,res_new,string_img
    if request.method == 'POST':
        if request.form.get('click') == 'Capture':
            print(333333333333)
            global capture
            capture=1

                
    return render_template('index.html')
# @app.route('/')
# def index():
#     full_filename = os.path.join(app.config['UPLOAD_FOLDER'], path)
#     print("full_filename.......................",full_filename)
#     return render_template("index.html", user_image = full_filename)

def parse_arguments(argv):
    parser = argparse.ArgumentParser()

    parser.add_argument('--debug', action='store_true',
                        help='Enable some debug outputs.')
    return parser.parse_args(argv)


if __name__ == '__main__':
    #main()
    app.run(host='192.168.0.158', port=8084)
